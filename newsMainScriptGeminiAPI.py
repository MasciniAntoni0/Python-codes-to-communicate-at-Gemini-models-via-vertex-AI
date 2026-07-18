"""
ai_news_report.py

Fetches the most recent AI news using Google's Gemini models on the Gemini
Developer API (the public, API-key-based "free tier" service — NOT Vertex
AI / Google Cloud), grounded with Google Search so the model can pull live
information from the web, then writes a NEUTRAL, cross-source summary of
each item into a Word (.docx) file, using the format:

    news title:
        news explanation

------------------------------------------------------------------------------
Authentication — Gemini API key (Google AI Studio / Gemini Developer API)
------------------------------------------------------------------------------
Unlike the Vertex AI variant of this script (which relies on Google Cloud
"Application Default Credentials" obtained via `gcloud auth
application-default login`), this version talks directly to the public
Gemini Developer API using a personal API key, which is what the Gemini
free tier uses.

1. Create a free API key at https://aistudio.google.com/apikey
2. Export it as an environment variable before running the script:

    export GEMINI_API_KEY="your-api-key-here"

(No Google Cloud project, billing account, or `gcloud` login is required
for the free tier — only the API key above.)

------------------------------------------------------------------------------
Requirements
------------------------------------------------------------------------------
    pip install google-genai python-docx

------------------------------------------------------------------------------
Run
------------------------------------------------------------------------------
    python ai_news_mainScript_GeminiAPIfree.py
"""
# Above: this is a module-level "docstring", i.e. a string literal delimited
# by triple quotes placed as the first statement of the file. Python
# automatically recognizes it and stores it in the module's special __doc__
# attribute; it serves as documentation readable both by a human (whoever
# opens the file) and by automated tools (e.g. help(), documentation
# generators). It has no effect when the script is executed: it is plain
# text, not executed code. In this variant of the file, the docstring has
# been updated to explain that authentication happens via a personal API key
# for the "Gemini Developer API" (the free tier), rather than through Google
# Cloud credentials (ADC) as used by the Vertex AI version.

# Empty line: stylistic separator for readability (PEP 8), has no effect.

from __future__ import annotations
# Special import (must be the first "executable" statement of the file).
# Enables the "postponed evaluation of annotations" feature (PEP 563): type
# annotations (e.g. "-> list[dict]") are treated as plain text and evaluated
# only if someone explicitly requests them at runtime, instead of being
# computed immediately. This allows using modern type syntax (like
# "list[dict]" or "int | None") even on older Python versions, without
# having to import the generic types from "typing".

# Empty separator line.

import json
# Imports the standard "json" module, which provides functions to convert
# (serialize) Python data structures into JSON-formatted text and to convert
# (deserialize) JSON text into Python data structures (dict, list, str, int,
# float, bool, None). Here it will be used to read the language model's
# text response as if it were a JSON object.

import os
# Imports the standard "os" module, which offers functions to interact with
# the operating system (environment variables, file paths, processes, etc.).
# In this version of the file it is particularly important: it will be used
# to read, via "os.environ.get(...)", the value of the GEMINI_API_KEY
# environment variable containing the personal API key needed to
# authenticate with the Gemini Developer API (the free tier).

import re
# Imports the standard "re" module (regular expressions): it allows
# searching, replacing, or extracting portions of text that match a given
# pattern. It will be used further down to "clean up" the text returned by
# the model from any Markdown code-block delimiters (```), before it can be
# parsed as JSON.

import sys
# Imports the standard "sys" module, which gives access to functionality
# tied to the Python interpreter itself: here it is used mainly to write
# messages to the standard error stream (sys.stderr) instead of the
# standard output stream, and to terminate the program with a specific exit
# code via sys.exit(...).

from datetime import datetime
# From the standard "datetime" module imports the "datetime" class, which
# represents a point in time (date + time). It will be used to obtain the
# current moment (datetime.now()) and to format it as text, for example to
# name the output file with today's date or to include it in the prompt
# sent to the model.

from google import genai
# Imports Google's "genai" package (Google Gen AI SDK), which provides the
# classes and functions needed to communicate with Gemini language models.
# The same SDK can talk either to Vertex AI (Google Cloud) or to the public
# Gemini Developer API (the one used here, via a plain API key): which of
# the two is used is determined by how the "client" object is created in
# the make_client() function below.

from google.genai import types
# From the same SDK imports the "types" sub-module, which contains support
# classes for building request configurations (e.g. which tools the model
# may use, the generation "temperature", the settings to force the use of a
# tool, etc.).

from google.genai.errors import APIError
# Imports the "APIError" exception class specific to the SDK: it will be
# used to intercept (via try/except) any errors returned by Google's remote
# service (e.g. exhausted quota, a model not available or not yet enabled
# for the API key in use, an invalid request), so the error can be handled
# and the script can move on to the next model in the fallback list, instead
# of the program crashing abruptly.

# Empty separator line.

from docx import Document
# From the external "python-docx" package imports the "Document" class,
# which represents an entire Word (.docx) document in memory: it allows
# adding headings, paragraphs, page breaks, bulleted lists, etc., and
# finally saving everything to disk as a .docx file.

from docx.shared import Pt
# From the same package imports "Pt" (Point), a class/unit of measurement
# used to express typographic dimensions (such as indents or spacing
# between paragraphs) in points, the standard unit used in publishing and
# word-processing programs (1 point = 1/72 of an inch).


# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------
# The lines above are a "block" comment: they only serve to visually
# separate the logical sections of the file for whoever reads the code;
# they have no effect on program execution.

GEMINI_API_KEY = "" 
# Defines a global variable (a constant by convention, written in
# UPPERCASE) that will hold the personal API key for the Gemini Developer
# API (the free tier). Unlike the Vertex AI version (which
# used a Google Cloud PROJECT_ID and ADC credentials), here the only secret
# needed is this API key, obtainable for free from Google AI Studio.
# So in this veriable You have to insert your own gemini API key,
# to use the gemini models via gemini API.

# Empty separator line.

# The base HTTP endpoint (URL) of the Gemini Developer API that every
# request will be sent to: this is the "REST" address of Google's servers
# that host the free-tier Gemini models.
GEMINI_API_BASE_URL = "https://generativelanguage.googleapis.com/v1beta"
# Defines a global variable (a constant by convention, written in
# UPPERCASE) holding the full base URL (Uniform Resource Locator, i.e. a
# web address) of the Gemini Developer API. Reading the URL piece by piece,
# from scratch, assuming nothing: "https://" is the PROTOCOL, the set of
# rules used to transfer data over the web, in its secure/encrypted variant
# (the "s" stands for the SSL/TLS layer that encrypts the traffic);
# "generativelanguage.googleapis.com" is the DOMAIN NAME, the
# human-readable address that identifies the server to contact — here
# "googleapis.com" is Google's family of API servers, and the
# "generativelanguage" prefix selects, within that family, the specific
# service that exposes the generative-language (Gemini) models; "/v1beta"
# is the first PATH segment and states the API VERSION being targeted —
# literally "version 1, beta channel", the version that currently exposes
# the Gemini models and their features (such as Google Search grounding).
# Everything that comes AFTER this base — the specific model name and the
# action to perform, for example
# ".../v1beta/models/gemini-3.1-flash-lite:generateContent" — is appended
# automatically by the SDK on each request, so it does not need to be
# written here. By defining this URL explicitly in one place (instead of
# relying on the SDK's hidden built-in default), the endpoint becomes
# visible, documented, and trivial to change should it ever move.

# Empty separator line.

# Ordered by priority — this is the Gemini API's free-tier model
# fallback chain: the SDK will try the first model, and only move on to
# the next one if the current model is unavailable, over quota, or errors
# out for any other reason.
# Comment: explains the logic behind the MODEL_FALLBACK_CHAIN list defined
# right below, i.e. that the models will be tried in sequence (from the
# first to the last) until one of them responds correctly; if a model is
# unavailable (for example because it is not yet enabled for the API key in
# use, or because the free quota assigned to it has been exhausted), the
# script automatically moves on to the next model in the list.

MODEL_FALLBACK_CHAIN = [
    # The first entry, "gemini-3.1-flash-lite", is the PRIMARY model — the
    # one tried first on every run, because it offers the most generous
    # free-tier quota (15 RPM, 250K TPM, 500 RPD) among the models below.
    "gemini-3.1-flash-lite",    # (15 RPM, 250K TPM, 500 RPD)  <- primary
    "gemini-2.5-flash-lite",
    "gemini-3.5-flash",
    "gemini-3-flash-preview",
    "gemini-2.5-flash",
]
# Creates a list (an ordered, mutable data structure) of strings, each one
# the name of a Gemini "flash" model (the lightest, fastest, and cheapest
# model family, suited to the API's free tier) available through the
# Gemini Developer API. The order matters and was chosen by PRIORITY: the
# program will first try "gemini-3.1-flash-lite" (the primary model, noted
# in the comment above as having the most generous free-tier quota limits:
# 15 requests per minute, 250 thousand tokens per minute, 500 requests per
# day); if this model is unavailable (not yet released for the key in use,
# quota exhausted, temporary service error, etc.), the script will try the
# following ones in order — "gemini-2.5-flash-lite", then
# "gemini-3.5-flash", then "gemini-3-flash-preview", finally
# "gemini-2.5-flash" — thereby implementing a "failover" mechanism
# (automatic fallback on error) identical in spirit to the one in the
# Vertex AI version, but applied here to a list of models specifically
# chosen for the API's free tier.

# Empty separator line.

NUM_ITEMS = 15            # how many news items to gather
# Defines an integer constant representing how many distinct news items the
# model must return in total. The comment alongside (after the #) clarifies
# the meaning of the value; the value itself will later be dynamically
# inserted into the prompt text sent to the model (see build_prompt).

INCLUDE_SOURCES = True    # append a "Sources" section (the grounded web links)
# Defines a boolean constant (True/False) acting as a switch ("feature
# flag"): if True, a section listing the web links used by the model as
# sources during its search ("grounding", i.e. anchoring the responses to
# real sources) will also be added to the final Word document.

OUTPUT_FILE = f"./File_name_{datetime.now():%Y-%m-%d}.docx"
# Builds the output file's path/name using an "f-string" (formatted
# string): everything between curly braces {} is evaluated as a Python
# expression and the result is inserted in its place in the final text.
# Here "datetime.now()" gets the current date and time, and ":%Y-%m-%d" is
# a formatting spec that turns it into a string like "2026-07-03". The
# final result is therefore a relative path such as
# "./File_name_2026-07-03.docx", in the current folder from which the script
# is launched.

TOPIC = "" 
# Is a variable that contains a string where the follow string is the topic that You want the news
# in the final word Document, so insert in the "" the new's topic.

NEWS_FEATURES = """"""
# This variable contain a docstring with the "features" about the news for example 
# if i want all the news about AI, I will insert in the docstring what kind of news I want about ai, so 
# new updates and models released by AI companies, AI regulaments by governments, futures updates/models released by AI companies and their features,
# hardware and more else. in this docsting You have to specific what kind of news you want obtain about the topic   

# -----------------------------------------------------------------------------
# Gemini client (API key via the Gemini Developer API — free tier)
# -----------------------------------------------------------------------------
# Section comment: visually separates the block of code that follows,
# dedicated to creating the client that talks to the model. Unlike the
# Vertex AI version, here the client does NOT authenticate via Google Cloud
# credentials (ADC), but via a simple API key.

def make_client() -> genai.Client:
    # Defines a function called "make_client", which takes no parameters
    # and, as indicated by the type annotation after "->", will return an
    # object of type genai.Client (thanks to "from __future__ import
    # annotations" this annotation is not evaluated at runtime, it remains
    # only documentation/for type-checkers like mypy).
    if not GEMINI_API_KEY:
        # Conditional "if" structure: the "not" operator applied to a
        # string returns True when the string is empty (or absent); so
        # this block runs only if the GEMINI_API_KEY environment variable
        # was not set at all (the fallback value "" read above with
        # os.environ.get).
        print(
            "[warn] GEMINI_API_KEY is not set. Get a free key at "
            "https://aistudio.google.com/apikey and export it as "
            "GEMINI_API_KEY.",
            file=sys.stderr,
        )
        # Calls the built-in "print" function to print a warning message.
        # Two adjacent strings are automatically concatenated by Python
        # into a single string. The "file=sys.stderr" parameter specifies
        # writing the message to the standard error stream (stderr)
        # instead of the standard one (stdout), by convention used for
        # diagnostic/log messages that are not part of the program's
        # "actual" output. The message guides the user to obtain a free
        # API key and export it as an environment variable before
        # re-running the script.
    # Passing api_key= (instead of vertexai=True) makes the SDK talk to the
    # public Gemini Developer API — the free-tier endpoint — using the
    # personal API key instead of Google Cloud ADC credentials.
    # Comment: explains the meaning of the api_key parameter in the call
    # below, i.e. that the SDK will use the provided key to authenticate
    # every request directly against the public Gemini Developer API (the
    # "free tier" service), instead of going through Vertex AI and Google
    # Cloud's ADC credentials as in the previous version of the script.
    # Build the "HTTP options" object that tells the SDK exactly which
    # server endpoint (URL) every request must be sent to. types.HttpOptions
    # is a small configuration container provided by the SDK; here two of
    # its fields are set:
    #   - base_url: the full base address of the API, taken from the
    #     GEMINI_API_BASE_URL constant defined in the Configuration section
    #     above (i.e. "https://generativelanguage.googleapis.com/v1beta").
    #   - api_version: deliberately set to an EMPTY string "". This needs an
    #     explanation from scratch: when api_version is left at its normal
    #     value, the SDK AUTOMATICALLY appends it as an extra path segment
    #     after base_url; since the version "v1beta" is ALREADY part of our
    #     base_url, letting the SDK add it again would produce a duplicated,
    #     invalid address like ".../v1beta/v1beta/models/...". Passing ""
    #     tells the SDK to add no extra version segment, so the final
    #     request URL stays exactly the one written in GEMINI_API_BASE_URL.
    http_options = types.HttpOptions(
        base_url=GEMINI_API_BASE_URL, api_version=""
    )
    # Creates the HttpOptions object described above and stores it in the
    # local variable "http_options", ready to be handed to the client below.
    return genai.Client(api_key=GEMINI_API_KEY, http_options=http_options)
    # "return" statement: creates an instance of the genai module's Client
    # class, passing as keyword arguments (1) the API key read above, which
    # authenticates every request, and (2) the http_options object just
    # built, which pins all requests to our explicit endpoint URL
    # (GEMINI_API_BASE_URL). By not passing "vertexai=True" (nor
    # "project"/"location"), the SDK defaults to using the public Gemini
    # Developer API instead of Vertex AI; returns that object to whoever
    # called the function, ending the function's execution.


def generate(client: genai.Client, prompt: str, *, use_search: bool,
             temperature: float = 0.3):
    # Defines the "generate" function, which receives: "client" (the object
    # created by make_client, annotated as type genai.Client), "prompt"
    # (the instruction text to send to the model, annotated as a string).
    # The bare asterisk "*" in the parameter list forces whoever calls the
    # function to specify all subsequent parameters by name (keyword-only),
    # not positionally: so "use_search" (boolean, required, indicates
    # whether to enable Google search) and "temperature" (decimal number,
    # with a default value of 0.3 if not specified, controlling how
    # "creative"/random the model's responses are: lower values = more
    # deterministic, factual responses).
    """Generate content with automatic failover across MODEL_FALLBACK_CHAIN.

    Returns (response, model_name).
    """
    # Function docstring: briefly describes what it does (generates content
    # by trying multiple models in sequence, following the priority order
    # of the MODEL_FALLBACK_CHAIN list) and what it returns (a tuple with
    # the response and the name of the model that actually answered). It is
    # documentation text, it produces no effect at runtime.
    tools = [types.Tool(google_search=types.GoogleSearch())] if use_search else None
    # Inline conditional expression ("expression if condition else
    # expression"): if use_search is true, builds a list containing a
    # single "Tool" object configured to use the "GoogleSearch" feature
    # (which allows the model to perform real web searches during
    # generation, i.e. "grounding"); otherwise the "tools" variable is set
    # to None, i.e. "no tools". This Google Search grounding feature is
    # also available through the Gemini Developer API with an API key (it
    # does not necessarily require Vertex AI), though it is subject to the
    # free-tier usage limits.
    # IMPORTANT — endpoint-specific behaviour (this is the bug fix): the
    # Vertex AI version of this script forced the model to call the search
    # tool by attaching a "tool_config" with function-calling mode "ANY".
    # On the Gemini Developer API (the free-tier endpoint used here) that
    # exact combination is REJECTED by the server with
    # "400 INVALID_ARGUMENT: Function calling config is set without
    # function_declarations." The reason, explained from scratch: mode "ANY"
    # means "you MUST call one of the functions I declared", but the built-in
    # google_search tool is NOT a user-declared function, and no
    # function_declarations are provided here — so the request is invalid.
    # Because that 400 is returned for every model that still has quota, the
    # fallback chain could never reach a working model (the newer models were
    # returning 429 "quota exhausted", while the models that DID still have
    # quota were being killed by this 400), making it look as though "all
    # models are out of daily requests". The correct approach on this
    # endpoint is to send NO tool_config and let grounding happen in the
    # default "AUTO" mode: for a news-gathering prompt the model still
    # grounds its answer with Google Search on its own.
    tool_config = None
    # Sets "tool_config" to None (Python's "no value") unconditionally, so no
    # function-calling configuration is attached to the request. This single
    # line is the fix: the request now carries only the google_search tool
    # (built above in "tools"), with no forced-invocation config to trigger
    # the 400 error, so grounding works and the model fallback chain can
    # actually fall through to a model that still has free-tier quota.
    config = types.GenerateContentConfig(
        temperature=temperature, tools=tools, tool_config=tool_config
    )
    # Creates a "GenerateContentConfig" object that groups together all the
    # generation request settings: the temperature received as a
    # parameter, the optional list of tools (tools), and the optional
    # configuration of those tools (tool_config) computed above. This
    # object will be passed to the SDK when the actual request is sent to
    # the model.

    # Empty separator line.

    last_err: Exception | None = None
    # Declares a "last_err" variable, with a type annotation indicating it
    # can hold either an instance of the base "Exception" class (an error)
    # or the special value "None" (no value/absence of error). It is
    # initialized to None: it will be used to "remember" the last error
    # encountered, in case all models fail, so it can be reported at the
    # end.
    for model in MODEL_FALLBACK_CHAIN:
        # "for" loop: iterates over the elements of the MODEL_FALLBACK_CHAIN
        # list one at a time, assigning each element (a string with the
        # model's name) to the "model" variable on each iteration,
        # following exactly the priority order established above (first
        # the primary model "gemini-3.1-flash-lite", then progressively the
        # others in case of unavailability).
        try:
            # "try" block: delimits a "risky" portion of code, i.e. one that
            # might raise an exception (error) during execution; if that
            # happens, execution jumps immediately to the corresponding
            # "except" block, instead of abruptly interrupting the entire
            # program.
            resp = client.models.generate_content(
                model=model, contents=prompt, config=config
            )
            # Invokes the client's "generate_content" method, i.e. performs
            # the actual network call to the Gemini Developer API, asking
            # the indicated model (the "model" string of the current
            # iteration) to generate a response to the "prompt" text,
            # according to the settings contained in "config". The result
            # (a response object) is saved into the local "resp" variable.
            if resp and resp.text:
                # Checks, using the logical "and" operator, that the "resp"
                # response object is not empty/false (e.g. None) AND that
                # its ".text" attribute is not empty: only if both
                # conditions are true is the request considered successful
                # with useful content.
                print(f"[ok] answered with {model}", file=sys.stderr)
                # Prints a log message to stderr confirming which model
                # actually produced the valid response, useful to
                # understand, in hindsight, whether the primary model was
                # used or one of the fallback models further down the
                # MODEL_FALLBACK_CHAIN list.
                return resp, model
                # "return" statement that returns a "tuple" (pair of
                # values) made of the response object and the name of the
                # model used; this immediately ends the function's
                # execution (and therefore also the for loop), because a
                # valid result has been obtained.
            last_err = RuntimeError("empty response")
            # If the "if" block above was not satisfied (empty response),
            # creates a generic "RuntimeError" exception with an
            # explanatory message and stores it in last_err, without
            # raising it immediately: the loop will continue trying the
            # next model.
        except APIError as exc:
            # "except" block: specifically intercepts exceptions of type
            # "APIError" (raised by the Google SDK in case of errors
            # returned by the remote service — for example a model not yet
            # available for the API key in use, free daily/per-minute quota
            # exhausted for that model, or an invalid request), assigning
            # the caught exception object to the local name "exc". It is
            # precisely this block that makes possible the "cycle through
            # the list" behavior required: whatever the reason a model in
            # the chain is unavailable, it is caught here, so the "for"
            # loop can continue and automatically try the next model in
            # MODEL_FALLBACK_CHAIN.
            last_err = exc
            # Saves the exception just caught into the last_err variable,
            # overwriting any previous value, so it can be reported later
            # if the subsequent models also fail.
            print(f"[warn] {model} failed: {exc}", file=sys.stderr)
            # Prints a warning message to stderr indicating which model
            # failed and why (the exception's text), for diagnostic
            # purposes; the "for" loop then naturally continues with the
            # next model in the fallback list.

    # Empty separator line (outside the for loop, but still inside the
    # "generate" function).

    raise RuntimeError(f"All models failed. Last error: {last_err}")
    # If the "for" loop ends naturally (i.e. no "return" was ever executed,
    # because ALL models in the chain failed or returned empty responses),
    # this "raise" statement deliberately raises a RuntimeError exception
    # with a message that includes the last recorded error, interrupting
    # execution and reporting the overall failure to whoever called this
    # function.


# -----------------------------------------------------------------------------
# Prompt
# -----------------------------------------------------------------------------
# Section comment: introduces the block of code that builds the instruction
# text (prompt) to send to the language model.

def build_prompt() -> str:
    # Defines the "build_prompt" function, with no parameters, which will
    # return (according to the "-> str" annotation) a string: the prompt
    # text.
    today = datetime.now().strftime("%Y-%m-%d")
    # Gets the current date and time with datetime.now() and applies the
    # ".strftime(...)" (string format time) method to convert it into a
    # text string in "YYYY-MM-DD" format (e.g. "2026-07-03"), which will be
    # inserted into the prompt text to give the model a time reference for
    # what counts as "recent".
    if not TOPIC:
        print("ERROR: Insert a topic to obtain the news.\n")
        sys.exit(1)
    # check that is inseted a topic to search the news
    else: 
        return f"""You are a neutral news analyst. Using Google Search, research the
MOST RECENT developments in {TOPIC} (roughly the last 2-4 weeks
relative to today, {today}).

Cover a broad range, including where relevant:
    {NEWS_FEATURES}

For EACH item, cross-check how MULTIPLE outlets report it, then write a NEUTRAL,
factual explanation that does not adopt any single publication's slant, framing,
or editorializing. Stick to verifiable facts and attribute claims where they are
contested.

Return ONLY a JSON array — no markdown fences, no commentary before or after —
of exactly {NUM_ITEMS} objects with this exact shape:

[
  {{"title": "<concise headline>",
    "explanation": "<neutral explanation, 3-6 sentences>"}}
]
"""
    # "return" statement returning a multi-line f-string (delimited by
    # triple quotes): all the English text is the actual instruction that
    # will be sent to the Gemini model. Inside it, "{today}" and
    # "{NUM_ITEMS}" are replaced with their respective values computed
    # above; the double curly braces "{{" and "}}" are a way to "escape"
    # braces inside f-strings, i.e. to make a single literal { or } brace
    # appear in the final text (needed to show the model an example of JSON
    # syntax, which uses braces to delimit objects) without Python
    # interpreting them as the start of an expression to evaluate.


# -----------------------------------------------------------------------------
# Parsing helpers
# -----------------------------------------------------------------------------
# Section comment: introduces the functions that process ("parse") the raw
# text returned by the model to turn it into structured data.

def parse_items(raw: str) -> list[dict]:
    # Defines the "parse_items" function, which receives a string "raw"
    # (the model's raw response text) and returns (per the annotation) a
    # list of dictionaries (list[dict]), i.e. a collection of news items
    # each represented as key-value pairs.
    """Robustly extract the JSON array from the model's text output."""
    # Docstring: explains that the function tries to "robustly" (tolerant
    # of small imperfections) extract the JSON array from the text produced
    # by the model, which sometimes may include unrequested Markdown
    # delimiters.
    text = raw.strip()
    # The ".strip()" method returns a copy of the "raw" string with leading
    # and trailing whitespace, tabs, and newlines removed; the result is
    # saved into a new local variable "text".
    text = re.sub(r"^```(?:json)?", "", text).strip()
    # "re.sub(pattern, replacement, text)" searches the text for the first
    # match of the given pattern and replaces it with the empty string. The
    # pattern, written as a "raw string" (r"..." to avoid special
    # interpretation of backslashes), is: "^" (start of string), followed
    # literally by three backticks, optionally followed (thanks to the
    # non-capturing group "(?:...)?" with "?") by the word "json". In
    # practice this removes any opening Markdown code-block delimiter such
    # as ```json or ``` placed at the start of the text (some models wrap
    # their output in these delimiters even when instructed not to). The
    # result is then cleaned again with ".strip()" of any residual
    # whitespace/newlines.
    text = re.sub(r"```$", "", text).strip()
    # Similarly, removes any closing Markdown block delimiter (three
    # backticks) placed at the end of the string (the "$" symbol in the
    # pattern means "end of string"), and cleans up any residual whitespace
    # again.

    # Empty separator line.

    start, end = text.find("["), text.rfind("]")
    # Multiple assignment: "text.find('[')" searches for the FIRST
    # occurrence of the '[' character in the text and returns its index
    # (position), or -1 if not found; "text.rfind(']')" instead searches
    # for the LAST occurrence of the ']' character (searching "from the
    # right"). The two indices are assigned respectively to "start" and
    # "end" in a single line via a tuple.
    if start != -1 and end != -1:
        # Checks that both delimiter characters were actually found in the
        # text (i.e. that neither index remained -1, the conventional value
        # meaning "not found").
        text = text[start : end + 1]
        # If both were found, "slices" the "text" string keeping only the
        # portion between index "start" (included) and "end + 1" (i.e. up
        # to and including the ']' character itself, since in Python
        # slicing the right end is exclusive). This isolates the JSON array
        # even if the model added introductory or trailing text around it.

    # Empty separator line.

    items = json.loads(text)
    # "json.loads" (load string) parses the "text" string as JSON-formatted
    # text and converts it into the corresponding Python data structure:
    # here a list of objects is expected, which in Python becomes a list of
    # dictionaries. If the text is not valid JSON, this call raises a
    # json.JSONDecodeError exception, handled by whoever calls this
    # function (see the main function further down).
    # keep only well-formed entries
    # Comment: explains the purpose of the following list comprehension,
    # i.e. filtering and normalizing only the well-formed entries,
    # discarding incomplete or malformed ones.
    return [
        {"title": str(it["title"]).strip(),
         "explanation": str(it["explanation"]).strip()}
        for it in items
        if isinstance(it, dict) and it.get("title") and it.get("explanation")
    ]
    # This is a "list comprehension": a compact way to build a new list by
    # iterating over an existing sequence. It reads from right to
    # left/from the center: "for it in items" iterates over each element of
    # the "items" list obtained from the JSON, calling it "it"; the final
    # "if" clause includes the element in the result only if: (1)
    # "isinstance(it, dict)" is true, i.e. the element is actually a
    # dictionary and not, for example, a bare string; (2) "it.get('title')"
    # returns a "truthy" value (not None, not an empty string — the .get
    # method returns None if the key does not exist, avoiding an error
    # compared to it["title"]); (3) likewise for "it.get('explanation')".
    # For each element that passes the filter, a new dictionary is built
    # with only the two keys "title" and "explanation", whose values are
    # forced to strings with "str(...)" (in case the model returned a
    # different type, e.g. a number) and cleaned of extra whitespace with
    # ".strip()". The final result is the "clean" list of dictionaries
    # returned by the function.


def extract_sources(resp) -> list[tuple[str, str]]:
    # Defines the "extract_sources" function, which receives the model's
    # response object ("resp", with no explicit type annotation) and
    # returns a list of tuples, each made of two strings (the title and URL
    # of a web source used by the model during its search).
    """Pull (title, url) pairs from the grounding metadata, de-duplicated."""
    # Docstring: explains that the function extracts (title, url) pairs
    # from the "grounding" metadata (i.e. the information about which web
    # pages were consulted by the model), removing any duplicates.
    seen: set[str] = set()
    # Creates an empty set, annotated as a set of strings. A "set" is an
    # unordered collection of unique elements: it will be used to keep
    # track of the URLs already encountered, allowing a very efficient
    # (constant-time) check of whether a URL has already been seen.
    out: list[tuple[str, str]] = []
    # Creates an empty list, annotated as a list of (string, string)
    # tuples: it will hold the final result, i.e. the unique (title, url)
    # pairs found, in the order they are encountered.
    try:
        # Opens a "try" block because accessing the nested attributes of
        # the response (candidates, grounding_metadata, etc.) might fail if
        # the structure is not the expected one, for example if the model
        # did not perform any search.
        chunks = resp.candidates[0].grounding_metadata.grounding_chunks or []
        # Accesses the "candidates" attribute of the response (the
        # possible response(s) generated by the model), takes the first
        # element (index [0], i.e. the main response), then its
        # "grounding_metadata" attribute (metadata about the anchoring to
        # web sources) and finally "grounding_chunks" (the individual
        # fragments/references to the web pages used). The "or []"
        # operator ensures that, if the obtained value were None or
        # otherwise "falsy" (e.g. an empty list), an empty literal list is
        # used instead, avoiding errors in the following "for" loop.
    except (AttributeError, IndexError, TypeError):
        # Catches three possible exception types that might occur during
        # the access above: "AttributeError" if an attribute (e.g.
        # grounding_metadata) does not exist on the object, "IndexError" if
        # the "candidates" list were empty (index [0] would not exist),
        # "TypeError" if an intermediate value were not of the expected
        # type and did not support the required operation.
        return out
        # In case of one of these errors, the function immediately returns
        # the "out" list as it is (empty, since it has not yet been filled
        # in), ending the function's execution without propagating the
        # exception to the caller.

    # Empty separator line (this code is reached only if the try block did
    # not raise any exceptions).

    for chunk in chunks:
        # "for" loop: iterates over each element of the "chunks" list
        # obtained above, assigning it in turn to the name "chunk".
        web = getattr(chunk, "web", None)
        # "getattr(object, attribute_name, default_value)" attempts to read
        # the "web" attribute from the "chunk" object; if that attribute
        # does not exist, it returns None instead of raising an
        # AttributeError exception. This is a "safe" way to access
        # potentially missing attributes.
        if web and getattr(web, "uri", None) and web.uri not in seen:
            # Checks, using the "and" operator (which evaluates conditions
            # from left to right and stops at the first false one), three
            # things in sequence: (1) that "web" is not None/falsy; (2)
            # that the "web" object has a non-empty "uri" attribute (the
            # page's address), read safely via getattr with a default of
            # None; (3) that this URI is not already present in the "seen"
            # set, i.e. that it has not already been encountered in a
            # previous iteration (to avoid duplicates in the final list).
            seen.add(web.uri)
            # If all conditions are true, adds the current URL to the
            # "seen" set, so that any future repetitions of the same URL
            # will be recognized and discarded by the check above.
            out.append((getattr(web, "title", None) or web.uri, web.uri))
            # Appends a new tuple (pair of values) to the "out" list: the
            # first element is the page's title, obtained safely via
            # getattr (default None if absent); if that title turned out
            # to be None or empty, the "or" operator makes it fall back to
            # using web.uri itself as the descriptive text; the second
            # element of the tuple is always the URL "web.uri".
    return out
    # At the end of the "for" loop (having examined all "chunk" elements),
    # returns the "out" list containing the unique (title, url) pairs
    # collected, in the order they were encountered in the response.


# -----------------------------------------------------------------------------
# Word output
# -----------------------------------------------------------------------------
# Section comment: introduces the function that generates the Word
# document.

def write_docx(items: list[dict], sources: list[tuple[str, str]], path: str) -> None:
    # Defines the "write_docx" function, which receives: "items" (the list
    # of title/explanation dictionaries produced by parse_items),
    # "sources" (the list of title/url pairs produced by extract_sources),
    # "path" (the string with the path of the file to create). The "->
    # None" annotation indicates that the function does not return any
    # useful value (its purpose is the side effect of writing a file to
    # disk).
    doc = Document()
    # Creates a new instance of the Document class (from python-docx), i.e.
    # a completely empty Word document, held in memory and not yet saved to
    # disk, to which content will gradually be added.

    # Empty separator line.

    doc.add_heading(f"Latest {TOPIC} News", level=0)
    # Calls the "add_heading" method on the document object to add a
    # heading with the text "Latest TOPIC News"; the "level=0" parameter
    # indicates that this is the highest-level heading (equivalent, in a
    # Word document, to the main "Title", distinct from level 1, 2, etc.
    # subheadings).
    stamp = doc.add_paragraph(
        f"Neutral cross-source summary — generated {datetime.now():%Y-%m-%d %H:%M}"
    )
    # Adds a new normal-text paragraph to the document, built as an
    # f-string: the fixed text "Neutral cross-source summary — generated "
    # is followed by the current date and time, this time formatted with
    # both the date and the time (":%Y-%m-%d %H:%M"). The "paragraph"
    # object returned by add_paragraph is saved into the "stamp" variable,
    # so it can be further modified in the next line.
    stamp.runs[0].italic = True
    # In python-docx, a paragraph's text is internally organized as a list
    # of "runs" (blocks of text with uniform formatting, accessible via the
    # ".runs" attribute); here the first (and in this case only) run of the
    # just-created paragraph is accessed, index [0], and its ".italic"
    # property is set to True, making that text appear in italics in the
    # final Word document.

    # Empty separator line.

    for item in items:
        # "for" loop: iterates over each dictionary "item" contained in the
        # "items" list received as a parameter, i.e. over each news item to
        # insert into the document, one at a time and in the order they
        # appear in the list.
        heading = doc.add_heading(level=2)
        # Adds a new empty level-2 heading to the document (a subheading,
        # smaller than the main title but larger than normal text); the
        # heading object created is saved into the local "heading" variable
        # so its text can be set right after (instead of passing the text
        # directly to add_heading, so as to control the formatting more
        # finely via a "run", as in the next line).
        heading.add_run(f"{item['title']}:")
        # Adds a text "run" to the just-created heading, i.e. a segment of
        # text with uniform formatting, whose content is the value
        # associated with the "title" key of the current "item" dictionary
        # (dictionary access with the syntax item['title']), followed by a
        # literal ":" character, according to the format required by the
        # file's initial docstring ("news title:").

        # Empty separator line.

        para = doc.add_paragraph(item["explanation"])
        # Adds a new normal-text paragraph to the document, whose content
        # is the value associated with the "explanation" key of the
        # current "item" dictionary (the neutral explanation of the news
        # item); the resulting paragraph object is saved into the local
        # "para" variable so its formatting can be modified right after.
        para.paragraph_format.left_indent = Pt(24)
        # Accesses the paragraph's formatting properties via the
        # ".paragraph_format" attribute and sets the left indent
        # ("left_indent", i.e. how far the text is shifted to the right
        # relative to the margin) to 24 typographic points (created with
        # the "Pt" class imported at the top of the file), so that the
        # explanation appears visually "indented" under its own heading,
        # reproducing the format described in the file's initial
        # docstring.
        para.paragraph_format.space_after = Pt(10)
        # Sets, again via the paragraph's formatting properties, the
        # vertical space to leave AFTER the paragraph itself (before the
        # following content) to 10 typographic points, to give a more
        # pleasant visual separation between one news item and the next.

    # Empty separator line (outside the for loop over the news items).

    if INCLUDE_SOURCES and sources:
        # Checks, using the "and" operator, two conditions: that the global
        # INCLUDE_SOURCES constant is True (the user wants the sources
        # section included) AND that the "sources" list received as a
        # parameter is not empty (i.e. there is actually at least one
        # source to show); only if both are true is the following indented
        # block executed.
        doc.add_page_break()
        # Inserts an explicit page break into the document, causing
        # everything that follows (the "Sources" section) to start on a
        # new page of the Word document, visually separating it from the
        # list of news items.
        doc.add_heading("Sources", level=1)
        # Adds a heading with the text "Sources", this time at level 1 (a
        # section heading, larger than a level-2 subheading but smaller
        # than the main level-0 title).
        for title, url in sources:
            # "for" loop with tuple "unpacking": on each iteration, the
            # current element of the "sources" list (a two-element tuple)
            # is automatically unpacked into the two local variables
            # "title" and "url".
            p = doc.add_paragraph(style="List Bullet")
            # Adds a new empty paragraph to the document, applying Word's
            # built-in "List Bullet" style (bulleted list), which will make
            # a bullet point appear before the text that will be added to
            # this paragraph. The paragraph object is saved into the local
            # variable "p".
            p.add_run(f"{title} — ").bold = True
            # Adds a first text "run" to paragraph "p" containing the
            # source's title followed by a space, an em dash "—" and
            # another space (f-string); the expression returned by
            # ".add_run(...)" is the newly created run object, on which the
            # ".bold" property is immediately set to True, rendering that
            # portion of text in bold in the final document.
            p.add_run(url)
            # Adds a second text "run" to the same paragraph "p", this time
            # containing simply the source's URL, with no further
            # formatting (not bold), which will therefore appear right
            # after the bold title, on the same line.

    # Empty separator line.

    doc.save(path)
    # Calls the document object's "save" method to physically write to
    # disk, at the path indicated by the "path" string (received as a
    # function parameter), all the content added so far in memory,
    # producing an actual .docx file that can be opened with Microsoft Word
    # or compatible programs.


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
# Section comment: introduces the main function that orchestrates the
# entire program flow, from start to finish.

def main() -> None:
    # Defines the "main" function, with no parameters, which does not
    # return any useful value ("-> None" annotation): it represents the
    # logical entry point of the entire script, which calls all the other
    # functions defined above in sequence.
    client = make_client()
    # Calls the previously defined "make_client" function to obtain a
    # client object already configured and authenticated against the
    # Gemini Developer API via the personal API key (free tier), saving it
    # into the local "client" variable.

    # Empty separator line.

    print(f"[info] gathering {TOPIC} news (grounded with Google Search)...",
          file=sys.stderr)
    # Prints an informational message to stderr signaling to the user that
    # the program is about to start gathering news via web search, to give
    # visual feedback while waiting (the call to the model can take a few
    # seconds).
    resp, _ = generate(client, build_prompt(), use_search=True)
    # Calls the previously defined "generate" function, passing the
    # client, the prompt text (obtained by calling "build_prompt()") and
    # the keyword parameter "use_search=True" to enable Google search. The
    # function returns a two-element tuple; the first (the actual
    # response) is assigned to the "resp" variable, while the second (the
    # name of the model used among those in MODEL_FALLBACK_CHAIN) is
    # assigned to the special name "_", which by convention in Python
    # indicates a value that is received but not needed for later use.

    # Empty separator line.

    try:
        # Opens a "try" block because interpreting the response text as
        # JSON (in the next line) might fail if the model returned text
        # that does not conform to expectations.
        items = parse_items(resp.text)
        # Calls the previously defined "parse_items" function, passing the
        # model's raw response text ("resp.text"), and saves the result
        # (the list of well-formed title/explanation dictionaries) into
        # the local "items" variable.
    except (json.JSONDecodeError, KeyError, TypeError) as exc:
        # Catches three possible error categories during parsing:
        # "json.JSONDecodeError" if the text is not valid JSON,
        # "KeyError" if an expected dictionary did not contain the
        # required keys (title/explanation), "TypeError" if a value had a
        # type incompatible with the required operations; the caught
        # exception is assigned to the local name "exc".
        print(f"[error] could not parse model output as JSON: {exc}",
              file=sys.stderr)
        # Prints an error message to stderr explaining that the model's
        # output could not be parsed as JSON, including the caught
        # exception's detail.
        print("----- raw output -----\n" + resp.text, file=sys.stderr)
        # Prints, for diagnostic purposes, the entire raw text returned by
        # the model (concatenated with the "+" operator to a descriptive
        # header and a newline "\n") to stderr, so the user can manually
        # inspect what the model actually generated and understand why
        # parsing failed.
        sys.exit(1)
        # Immediately terminates the entire program's execution with exit
        # code 1, the value conventionally used by operating systems to
        # indicate that the program ended due to an error (0 would instead
        # indicate termination without errors).

    # Empty separator line (reached only if parsing succeeded).

    if not items:
        # Checks whether the "items" list is empty: the "not" operator
        # applied to a list returns True if the list contains no elements
        # (an empty list is considered "falsy" in a boolean context in
        # Python).
        print("[error] no news items were produced.", file=sys.stderr)
        # If no valid news item was produced (even if the JSON was perhaps
        # syntactically correct but empty, or all elements were discarded
        # by the filter in parse_items), prints an error message to stderr
        # signaling this.
        sys.exit(1)
        # Terminates the program with exit code 1, indicating an error, for
        # the same reason explained above.

    # Empty separator line.

    sources = extract_sources(resp)
    # Calls the previously defined "extract_sources" function, passing the
    # full response object "resp" (which also contains the grounding
    # metadata), and saves the result (the list of unique title/url pairs
    # of the web sources) into the local "sources" variable.
    print(f"[info] {len(sources)} source links used by the model:", file=sys.stderr)
    # Prints an informational message to stderr indicating how many
    # sources (computed with the built-in "len" function, which returns the
    # number of elements in a list) were actually used by the model.
    for title, url in sources:
        # "for" loop with tuple unpacking, similar to the one seen in
        # write_docx: iterates over each (title, url) pair in the
        # "sources" list.
        print(f"    - {title}: {url}", file=sys.stderr)
        # For each source, prints a line to stderr formatted with four
        # spaces of indentation, a dash, the title, a colon, and the URL,
        # to show the user a readable list of sources directly in the
        # terminal, in addition to those that will be written into the
        # Word document.

    # Empty separator line.

    write_docx(items, sources, OUTPUT_FILE)
    # Calls the previously defined "write_docx" function, passing the list
    # of news items, the list of sources, and the output path computed at
    # the top of the file (the global OUTPUT_FILE constant), to physically
    # generate the final Word document on disk.

    # Empty separator line.

    print(f"[done] wrote {len(items)} items to {OUTPUT_FILE}")
    # Prints (this time to standard stdout, not stderr, since the "file"
    # parameter is not specified, which defaults to sys.stdout) a
    # completion message confirming how many news items were written and
    # to which file, representing the program's "final", user-facing
    # output.
    if INCLUDE_SOURCES:
        # Checks the global INCLUDE_SOURCES constant again: if true, also
        # prints a message about the included sources.
        print(f"[done] included {len(sources)} sources")
        # Prints a message to stdout confirming how many sources were
        # included in the final document.


if __name__ == "__main__":
    # This is a standard Python idiom: the special "__name__" variable
    # automatically equals the string "__main__" when the file is run
    # directly (e.g. with "python ai_news_mainScript_GeminiAPIfree.py"),
    # while it takes on the module's own name if the file is instead
    # imported from another Python script (with "import
    # ai_news_mainScript_GeminiAPIfree"). This check therefore allows the
    # following block to run only when the script is launched directly,
    # and not when its functions are reused elsewhere via import.
    main()
    # Calls the previously defined "main" function, actually starting the
    # entire program flow: creating the client authenticated via API key,
    # building the prompt, calling the model (with automatic failover
    # through MODEL_FALLBACK_CHAIN), parsing the response, extracting the
    # sources, and writing the final Word document.
