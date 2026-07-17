"""
ai_news_report.py

Fetches the most recent news about a specific topic specified in the promt's function using Google's Gemini models on Vertex AI,
grounded with Google Search so the model can pull live information from the web,
then writes a NEUTRAL, cross-source summary of each item into a Word (.docx)
file, using the format:

    news title:
        news explanation

------------------------------------------------------------------------------
Authentication — Google Cloud Application Default Credentials (ADC)
------------------------------------------------------------------------------
Run once on your machine:

    gcloud auth application-default login
    gcloud config set project YOUR_PROJECT_ID

Then either edit PROJECT_ID below or export it:

    export GOOGLE_CLOUD_PROJECT=your-gcp-project-id

------------------------------------------------------------------------------
Requirements
------------------------------------------------------------------------------
    pip install google-genai python-docx

------------------------------------------------------------------------------
Run
------------------------------------------------------------------------------
    python news_mainScript.py or python3 news_mainScript.py
"""
# Above: this is a module "docstring", i.e. a string literal delimited by
# triple quotes placed as the first statement of the file. Python recognizes
# it automatically and stores it in the module's special __doc__ attribute; it
# serves as documentation readable both by a human (whoever opens the file)
# and by automated tools (e.g. help(), documentation generators). It has no
# effect when the script is executed: it is plain text, not executed code.

# Blank line: stylistic separator for readability (PEP 8), has no effect.

from __future__ import annotations
# Special import (must be the first "executable" statement of the file).
# Enables the "postponed evaluation of annotations" feature (PEP 563):
# type annotations (e.g. "-> list[dict]") are treated as plain
# text and evaluated only if someone explicitly requests them at runtime,
# instead of being computed immediately. This allows using modern type
# syntax (like "list[dict]" or "int | None") even on older
# Python versions, without having to import generic types from "typing".

# Separator blank line.

import json
# Imports the standard "json" module, which provides functions to convert
# (serialize) Python data structures into JSON-format text and to
# convert (deserialize) JSON text into Python data structures
# (dict, list, str, int, float, bool, None). Here it will be used to read the
# text response of the language model as if it were a JSON object.

import os
# Imports the standard "os" module, which offers functions to interact with the
# operating system (environment variables, file paths, processes, etc.).
# It is imported to make available any system operations related to
# configuration (for example reading environment variables),
# consistent with what is described in the initial docstring.

import re
# Imports the standard "re" module (regular expressions): it allows
# searching, replacing, or extracting portions of text that match a
# given pattern. It will be used further on to
# "clean up" the text returned by the model from any Markdown code
# block delimiters (```), before it can be parsed as JSON.

import sys
# Imports the standard "sys" module, which gives access to functionality related
# to the Python interpreter itself: here it is mainly used to write
# messages to the standard error stream (sys.stderr) instead of the
# standard output stream, and to terminate the program with a specific
# exit code via sys.exit(...).

from datetime import datetime
# From the standard "datetime" module imports the "datetime" class, which
# represents a point in time (date + time). It will be used to obtain the
# current moment (datetime.now()) and to format it as text, for example
# to name the output file with today's date or to state it in the
# prompt sent to the model.

from google import genai
# Imports Google's "genai" package (Google Gen AI SDK), which provides the
# classes and functions needed to communicate with the Gemini
# language models, either via Google's public API or via Vertex AI (Google
# Cloud's artificial intelligence platform). From this module
# the "client" object used to send requests to the model will be obtained.

from google.genai import types
# From the same SDK imports the "types" sub-module, which contains support
# classes for building request configurations (for example which
# tools the model can use, the generation "temperature", the
# settings to force the use of a tool, etc.).

from google.genai.errors import APIError
# Imports the SDK-specific "APIError" exception class: it will be used
# to catch (via try/except) any errors returned by Google's
# remote service (for example quota exhausted, model not
# available, invalid request), so that the error can be handled instead
# of letting the program terminate abruptly.

# Separator blank line.

from docx import Document
# From the external "python-docx" package imports the "Document" class, which
# represents an entire Word (.docx) document in memory: it allows
# adding headings, paragraphs, page breaks, bullet lists, etc.
# and finally saving everything to disk as a .docx file.

from docx.shared import Pt
# From the same package imports "Pt" (Point), a class/unit of measure
# used to express typographic sizes (such as indents or spacing between
# paragraphs) in points, the standard unit used in publishing and word
# processing programs (1 point = 1/72 of an inch).


# -----------------------------------------------------------------------------
# Configuration
# -----------------------------------------------------------------------------
# The lines above are a "block" comment: they only serve to visually
# separate the logical sections of the file for whoever reads the code; they have
# no effect on the execution of the program.

TOPIC = "" 
# Is a variable that contains a string where the follow string is the topic that You want the news
# in the final word Document, so insert in the "" the new's topic.

NEWS_FEATURES = """"""
# This variable contain a docstring with the "features" about the news for example 
# if i want all the news about AI, I will insert in the docstring what kind of news I want about ai, so 
# new updates and models released by AI companies, AI regulaments by governments, futures updates/models released by AI companies and their features,
# hardware and more else. in this docsting You have to specific what kind of news you want obtain about the topic   
PROJECT_ID = "your-gcp-project-id"
# Defines a global variable (constant by convention, written in
# UPPERCASE) containing the identifier of the Google Cloud project to use
# to authenticate and bill calls to Vertex AI. This value must
# correspond to a real project on which the ADC login
# (Application Default Credentials) described in the initial docstring has been performed.

# "global" works for the newest 2.5 models; "us-central1" is a safe alternative.
# Explanatory comment: clarifies the possible alternatives for the
# next variable (LOCATION), explaining in which cases it is convenient to use one value or
# the other. It is not executed code.

LOCATION = ""
# Defines the geographic/infrastructural region of Google Cloud
# (data center) to which requests to the Vertex AI model will be
# routed. Some models or features are only available in certain
# regions, so this value must be compatible with the models listed below.
# For example "global", "us-central1" and more else.

# Separator blank line.

# Tried in order — automatic failover if one is unavailable / quota-limited.
# Comment: explains the logic behind the MODELS list defined right below,
# i.e. that they will be tried in sequence until one responds correctly.

MODELS = ["gemini-2.5-pro", "gemini-2.5-flash"]
# Creates a list (ordered, mutable data structure) of strings, each
# the name of a Gemini model available on Vertex AI. Order matters: the
# program will first try "gemini-2.5-pro" (more capable but possibly
# slower/more expensive or subject to quota limits) and, only if it fails, will move to the
# next one "gemini-2.5-flash" (faster/cheaper), thus implementing a
# "failover" mechanism (automatic fallback in case of error).

# Separator blank line.

NUM_ITEMS = 15            # how many news items to gather Inserted 15 as default but You can edit it and 
# Defines an integer constant representing how many distinct news items the
# model must return in total. The comment alongside (after the #)
# clarifies the meaning of the value; the value itself will then be inserted
# dynamically into the prompt text sent to the model (see build_prompt).

INCLUDE_SOURCES = True    # append a "Sources" section (the grounded web links)
# Defines a boolean constant (True/False) that acts as a
# switch ("feature flag"): if True, a section with the list of
# web links used by the model as sources during the
# search ("grounding", i.e. anchoring responses to real sources) will also be added to the final Word document.

OUTPUT_FILE = f"File_name{datetime.now():%Y-%m-%d}.docx"   # Substitute "File_name" with the name that you want save the final file 
# Builds the output file path/name using an "f-string"
# (formatted string): everything between curly braces {} is
# evaluated as a Python expression and the result is inserted in its
# place in the final text. Here "datetime.now()" obtains the current
# date and time, and ":%Y-%m-%d" is a formatting specifier that turns it
# into a string like "2026-07-03". The final result is therefore a
# relative path like "./File_name_2026-07-03.docx", in the
# current folder from which the script is launched.


# -----------------------------------------------------------------------------
# Gemini client (ADC via Vertex AI)
# -----------------------------------------------------------------------------
# Section comment: visually separates the following block of code,
# dedicated to creating the client that talks to the model.

def make_client() -> genai.Client:
    # Defines a function called "make_client", which receives no parameters
    # and, as indicated by the type annotation after "->", will return an
    # object of type genai.Client (thanks to "from __future__ import
    # annotations" this annotation is not evaluated at runtime, it remains
    # only documentation/for type-checkers like mypy).
    if PROJECT_ID == "your-gcp-project-id":
        # "if" conditional structure: checks whether the PROJECT_ID constant is
        # still set to the example placeholder value (not yet
        # customized by the user). The "==" operator compares two
        # strings for content equality.
        print(
            "[warn] PROJECT_ID is not set. Edit PROJECT_ID or export "
            "GOOGLE_CLOUD_PROJECT.",
            file=sys.stderr,
        )
        # Calls the built-in function "print" to print a warning
        # message. Two adjacent strings are automatically concatenated
        # by Python into a single string. The "file=sys.stderr" parameter
        # indicates to write the message to the standard error stream
        # (stderr) instead of the standard output stream (stdout), a convention
        # used for diagnostic/log messages that are not part of
        # the program's "actual" output.
    # vertexai=True makes the SDK authenticate through ADC automatically.
    # Comment: explains the meaning of the vertexai=True parameter in the
    # call below, i.e. that the SDK will automatically use the
    # application's default credentials (ADC) configured via
    # "gcloud auth application-default login", without needing to
    # manually manage API keys.
    return genai.Client(vertexai=True, project=PROJECT_ID, location=LOCATION)
    # "return" statement: creates an instance of the genai module's Client
    # class, passing as keyword arguments the flag to enable
    # Vertex AI mode, the project identifier and the region
    # defined above as global constants; returns that object to whoever
    # called the function, terminating the function's execution.


def generate(client: genai.Client, prompt: str, *, use_search: bool,
             temperature: float = 0.3):
    # Defines the "generate" function, which receives: "client" (the object
    # created by make_client, annotated as type genai.Client), "prompt" (the
    # instruction text to send to the model, annotated as a string).
    # The lone asterisk "*" in the parameter list forces whoever calls
    # the function to specify all subsequent parameters by name
    # (keyword-only), not by position: so "use_search" (boolean,
    # required, indicates whether to enable Google search) and "temperature"
    # (decimal number, with default value 0.3 if not specified, which
    # controls how "creative"/random the model's responses are: lower
    # values = more deterministic and factual responses).
    """Generate content with automatic failover across MODELS.

    Returns (response, model_name).
    """
    # Function docstring: briefly describes what it does (generates
    # content trying multiple models in sequence) and what it returns (a
    # tuple with the response and the name of the model that actually
    # responded). It is documentation text, produces no effect at runtime.
    tools = [types.Tool(google_search=types.GoogleSearch())] if use_search else None
    # Inline conditional expression ("expression if condition else
    # expression"): if use_search is true, builds a list containing a
    # single "Tool" object configured to use the
    # "GoogleSearch" feature (which allows the model to perform real
    # web searches during generation, i.e. "grounding"); otherwise the
    # "tools" variable is set to None, i.e. "no tool".
    # Force the model to actually invoke the Google Search tool instead of
    # letting it decide (AUTO) whether grounding is "needed".
    # Comment: explains the reason for the following block, i.e. explicitly
    # forcing the use of the search tool instead of leaving it to the
    # model to autonomously decide (default behavior, called
    # "AUTO") whether or not to perform a web search.
    tool_config = (
        types.ToolConfig(
            function_calling_config=types.FunctionCallingConfig(mode="ANY")
        )
        if use_search
        else None
    )
    # Similar to the line above, with a conditional expression
    # written across multiple lines for readability: if use_search is true, creates a
    # ToolConfig object whose "function calling" configuration (i.e. the
    # mechanism by which the model decides to invoke external tools) has
    # mode "ANY", which forces the model to invoke at least one
    # available tool (here, the search) instead of letting it decide freely;
    # if use_search is false, tool_config stays None (no configuration).
    config = types.GenerateContentConfig(
        temperature=temperature, tools=tools, tool_config=tool_config
    )
    # Creates a "GenerateContentConfig" object that groups together all the
    # settings for the generation request: the temperature received
    # as a parameter, the optional list of tools (tools), and the optional
    # tool configuration (tool_config) computed above. This
    # object will be passed to the SDK when the actual
    # request is sent to the model.

    # Separator blank line.

    last_err: Exception | None = None
    # Declares a variable "last_err", with a type annotation indicating
    # that it can contain either an instance of the base class "Exception" (error)
    # or the special value "None" (no value/absence of error).
    # It is initialized to None: it will be used to "remember" the last
    # error encountered, in case all models fail, so it can be
    # reported at the end.
    for model in MODELS:
        # "for" loop: iterates over the elements of the MODELS list one at
        # a time, assigning each element (a string with the model name)
        # to the "model" variable on each iteration, in the order of
        # preference established above (first "pro", then "flash").
        try:
            # "try" block: delimits a portion of "at-risk" code,
            # i.e. that could raise an exception (error) during
            # execution; if that happens, execution jumps immediately
            # to the corresponding "except" block, instead of abruptly
            # interrupting the whole program.
            resp = client.models.generate_content(
                model=model, contents=prompt, config=config
            )
            # Invokes the client's "generate_content" method, i.e. performs
            # the actual network call to the Vertex AI service,
            # asking the indicated model (the "model" string of the
            # current iteration) to generate a response to the
            # "prompt" text, according to the settings contained in "config". The
            # result (a response object) is saved in the local
            # variable "resp".
            if resp and resp.text:
                # Checks, with the logical operator "and", that the response
                # object "resp" is not empty/false (e.g. None) AND that
                # its text attribute ".text" is not empty: only if
                # both conditions are true is the request considered
                # successful with usable content.
                print(f"[ok] answered with {model}", file=sys.stderr)
                # Prints a log message to stderr confirming which
                # model actually produced the valid response,
                # useful for figuring out, after the fact, whether the
                # main model or the fallback one was used.
                return resp, model
                # "return" statement that returns a "tuple" (pair of
                # values) made up of the response object and the name of the
                # model used; this immediately stops
                # the function's execution (and therefore also the for
                # loop), because a valid result has been obtained.
            last_err = RuntimeError("empty response")
            # If the "if" block above was not satisfied (empty
            # response), creates a generic "RuntimeError" exception with
            # an explanatory message and stores it in last_err, without
            # however raising it right away: it will keep trying with the
            # next model in the loop.
        except APIError as exc:
            # "except" block: specifically catches exceptions of
            # type "APIError" (raised by the Google SDK in case of errors
            # returned by the remote service), assigning the caught
            # exception object to the local name "exc".
            last_err = exc
            # Saves the just-caught exception into the
            # last_err variable, overwriting any previous value, so
            # it can be reported later if subsequent models
            # should also fail.
            print(f"[warn] {model} failed: {exc}", file=sys.stderr)
            # Prints a warning message to stderr indicating which
            # model failed and why (the exception
            # text), for diagnostic purposes; then the "for" loop
            # naturally continues with the next model in the list.

    # Separator blank line (outside the for loop, but still inside the
    # "generate" function).

    raise RuntimeError(f"All models failed. Last error: {last_err}")
    # If the "for" loop ends naturally (i.e. no "return" was ever
    # executed, because all models failed or returned empty
    # responses), this "raise" statement deliberately raises
    # a RuntimeError exception with a message that includes the last
    # recorded error, interrupting execution and reporting the
    # overall failure to whoever called this function.


# -----------------------------------------------------------------------------
# Prompt
# -----------------------------------------------------------------------------
# Section comment: introduces the block of code that builds the
# instruction text (prompt) to send to the language model.

def build_prompt() -> str:
    # Defines the "build_prompt" function, with no parameters, which will return
    # (according to the "-> str" annotation) a string: the prompt text.
    today = datetime.now().strftime("%Y-%m-%d")
    # Gets the current date and time with datetime.now() and applies the
    # ".strftime(...)" (string format time) method to convert them into a
    # text string in the format "YYYY-MM-DD" (e.g. "2026-07-03"), which will be
    # inserted into the prompt text to give the model a
    # temporal reference against which to define what is "recent".
    if len(TOPIC) == 0:
        sys.exit("ERROR: Insert a valid topic to continue. Try again, insert in the Variable TOPIC a string with the topic's name.")
    # Is a check to verify that was insered a topic to take the news
    # 
    else:
        return f"""You are a neutral news analyst. Using Google Search, research the
MOST RECENT news about {TOPIC} (roughly the last 2-4 weeks
relative to today, {today}).

Cover a broad range, including where relevant: 
    {NEWS_FEATURES}

For EACH item, cross-check how MULTIPLE outlets report it, then write a NEUTRAL,
factual explanation that does not adopt any single publication's slant, framing,
or editorializing. Stick to verifiable facts and attribute claims where they are
contested.

Return ONLY a JSON array — no markdown fences, no commentary before or after —
of exactly {NUM_ITEMS} objects with this exact shape:

[
  {{"title": "<concise headline>",
    "explanation": "<neutral explanation, 3-6 sentences>"}}
]
"""
    # "return" statement that returns a multi-line f-string (delimited
    # by triple quotes): all the English text is the actual
    # instruction that will be sent to the Gemini model. Within it, "{today}"
    # and "{NUM_ITEMS}" are replaced with their respective values computed
    # above; the double curly braces "{{" and "}}" are a way to
    # "escape" the braces in f-strings, i.e. to make a single literal
    # brace { or } appear in the final text (needed to show
    # the model an example of JSON syntax, which uses braces to
    # delimit objects) without Python interpreting them as the start of
    # an expression to evaluate.


# -----------------------------------------------------------------------------
# Parsing helpers
# -----------------------------------------------------------------------------
# Section comment: introduces the functions that process ("parse") the
# raw text returned by the model to turn it into structured data.

def parse_items(raw: str) -> list[dict]:
    # Defines the "parse_items" function, which receives a string "raw" (the
    # raw response text from the model) and returns (per annotation)
    # a list of dictionaries (list[dict]), i.e. a collection of news items
    # each represented as key-value pairs.
    """Robustly extract the JSON array from the model's text output."""
    # Docstring: explains that the function tries to "robustly"
    # (tolerant of small imperfections) extract the JSON array from the text produced by the
    # model, which sometimes may include unrequested Markdown delimiters.
    text = raw.strip()
    # The ".strip()" method returns a copy of the "raw" string with leading and trailing
    # whitespace, tabs, and newlines removed; the
    # result is saved into a new local variable "text".
    text = re.sub(r"^```(?:json)?", "", text).strip()
    # "re.sub(pattern, replacement, text)" searches the text for the first
    # match of the given pattern and replaces it with the empty
    # string. The pattern, written as a "raw string" (r"..." to avoid
    # special interpretation of backslashes), is: "^" (start of
    # string), followed literally by three backticks, optionally followed
    # (because of the non-capturing group "(?:...)?" with "?")
    # by the word "json". In practice it removes any Markdown code
    # block opening like ```json or ``` placed at the start of the
    # text (some models wrap the output in these delimiters even
    # if instructed not to). The result is then cleaned up again with
    # ".strip()" of any remaining spaces/newlines.
    text = re.sub(r"```$", "", text).strip()
    # Similarly, removes any Markdown block closing (three
    # backticks) placed at the end of the string (the "$" symbol in the pattern
    # means "end of string"), and cleans up remaining spaces again.

    # Separator blank line.

    start, end = text.find("["), text.rfind("]")
    # Multiple assignment: "text.find('[')" searches for the FIRST
    # occurrence of the character '[' in the text and returns its index (position), or -1
    # if not found; "text.rfind(']')" instead searches for the LAST
    # occurrence of the character ']' (search "from the right"). The two
    # indices are assigned respectively to "start" and "end" in a single line via a tuple.
    if start != -1 and end != -1:
        # Checks that both delimiter characters were
        # actually found in the text (i.e. that neither index
        # remained -1, the conventional value indicating "not found").
        text = text[start : end + 1]
        # If both were found, "slices" the "text" string,
        # keeping only the portion between index
        # "start" (included) and "end + 1" (i.e. up to and including the ']'
        # character itself, since in Python slicing the right
        # endpoint is excluded). This isolates the JSON array even if
        # the model added introductory or trailing text
        # around it.

    # Separator blank line.

    items = json.loads(text)
    # "json.loads" (load string) parses the "text" string as
    # JSON-format text and converts it into the corresponding Python data
    # structure: a list of objects is expected here, which in Python becomes a
    # list of dictionaries. If the text is not valid JSON, this call
    # raises a json.JSONDecodeError exception, handled by whoever calls
    # this function (see the main function further below).
    # keep only well-formed entries
    # Comment: explains the purpose of the following list comprehension, i.e.
    # filter and normalize only the well-formed entries, discarding those
    # that are incomplete or malformed.
    return [
        {"title": str(it["title"]).strip(),
         "explanation": str(it["explanation"]).strip()}
        for it in items
        if isinstance(it, dict) and it.get("title") and it.get("explanation")
    ]
    # This is a "list comprehension": a compact way to build a
    # new list by iterating over an existing sequence. It reads from right to
    # left/from the middle: "for it in items" iterates over each element of the
    # "items" list obtained from the JSON, calling it "it"; the final "if"
    # clause includes the element in the result only if: (1) "isinstance(it,
    # dict)" is true, i.e. the element is actually a dictionary and not,
    # for example, an isolated string; (2) "it.get('title')" returns a
    # "truthy" value (not None, not an empty string — the .get method
    # returns None if the key doesn't exist, avoiding an error compared
    # to it["title"]); (3) likewise for "it.get('explanation')". For
    # each element that passes the filter, a new
    # dictionary is built with only the two keys "title" and "explanation", whose
    # values are coerced to string with "str(...)" (in case the model
    # had returned a different type, e.g. a number) and cleaned of extra
    # whitespace with ".strip()". The final result is the list of
    # "cleaned" dictionaries returned by the function.


def extract_sources(resp) -> list[tuple[str, str]]:
    # Defines the "extract_sources" function, which receives the model's
    # response object ("resp", with no explicit type annotation) and returns
    # a list of tuples, each made up of two strings (title and URL of
    # a web source used by the model during the search).
    """Pull (title, url) pairs from the grounding metadata, de-duplicated."""
    # Docstring: explains that the function extracts (title, url) pairs from the
    # "grounding" metadata (i.e. information about which web pages
    # were consulted by the model), removing any duplicates.
    seen: set[str] = set()
    # Creates an empty set, annotated as a set of strings. A
    # "set" is an unordered collection of unique elements: it will be used to
    # keep track of URLs already encountered, allowing very
    # efficient (constant-time) checking of whether a URL has already been seen.
    out: list[tuple[str, str]] = []
    # Creates an empty list, annotated as a list of tuples (string, string):
    # it will hold the final result, i.e. the unique (title, url) pairs
    # found, in the order they are encountered.
    try:
        # Opens a "try" block because accessing the nested
        # attributes of the response (candidates, grounding_metadata, etc.)
        # could fail if the structure is not the expected one, for example if the
        # model did not perform any search.
        chunks = resp.candidates[0].grounding_metadata.grounding_chunks or []
        # Accesses the "candidates" attribute of the response (the possible
        # response(s) generated by the model), takes the first element (index
        # [0], i.e. the main response), then its "grounding_metadata" attribute
        # (metadata about grounding to web sources) and
        # finally "grounding_chunks" (the individual fragments/references to
        # the web pages used). The "or []" operator ensures that, if the
        # obtained value were None or otherwise "falsy" (for example an
        # empty list), an empty literal list is used in its place,
        # avoiding errors in the subsequent "for" loop.
    except (AttributeError, IndexError, TypeError):
        # Catches three possible types of exception that could
        # occur during the access above: "AttributeError" if an
        # attribute (e.g. grounding_metadata) doesn't exist on the object,
        # "IndexError" if the "candidates" list were empty (index [0]
        # wouldn't exist), "TypeError" if an intermediate value were not
        # of the expected type and didn't support the required operation.
        return out
        # In case of one of these errors, the function immediately
        # returns the "out" list as is (empty, since it hasn't
        # yet been filled), interrupting the function's execution
        # without propagating the exception to the caller.

    # Separator blank line (this code is reached only if the try
    # block did not raise exceptions).

    for chunk in chunks:
        # "for" loop: iterates over each element of the "chunks" list
        # obtained above, assigning it in turn to the name "chunk".
        web = getattr(chunk, "web", None)
        # "getattr(object, attribute_name, default_value)" tries to
        # read the "web" attribute from the "chunk" object; if that attribute
        # doesn't exist, it returns None instead of raising an
        # AttributeError exception. This is a "safe" way of accessing
        # potentially missing attributes.
        if web and getattr(web, "uri", None) and web.uri not in seen:
            # Checks, with the "and" operator (which evaluates conditions
            # left to right and stops at the first false one), three things in
            # sequence: (1) that "web" is not None/falsy; (2) that the
            # "web" object has a "uri" attribute (the page's address) that is not
            # empty, again read safely with getattr and default None;
            # (3) that this URI is not already present in the "seen" set,
            # i.e. that it wasn't already encountered in a
            # previous iteration (to avoid duplicates in the final list).
            seen.add(web.uri)
            # If all conditions are true, adds the current URL
            # to the "seen" set, so that any future repetitions of the
            # same URL are recognized and discarded by the check above.
            out.append((getattr(web, "title", None) or web.uri, web.uri))
            # Adds a new tuple (pair of values) to the "out" list:
            # the first element is the page title, obtained safely with
            # getattr (default None if absent); if that
            # title turns out to be None or empty, the "or" operator makes it so
            # the same "web.uri" is used in its place as a fallback
            # descriptive text; the second element of the tuple is
            # always the URL "web.uri".
    return out
    # At the end of the "for" loop (which examined all the "chunk"s),
    # returns the "out" list containing the unique (title, url) pairs
    # collected, in the order they were encountered in the response.


# -----------------------------------------------------------------------------
# Word output
# -----------------------------------------------------------------------------
# Section comment: introduces the function that generates the Word document.

def write_docx(items: list[dict], sources: list[tuple[str, str]], path: str) -> None:
    # Defines the "write_docx" function, which receives: "items" (the list of
    # title/explanation dictionaries produced by parse_items), "sources" (the
    # list of title/url pairs produced by extract_sources), "path" (the
    # string with the path of the file to create). The "-> None" annotation
    # indicates that the function does not return any useful value (its
    # purpose is the side effect of writing a file to disk).
    doc = Document()
    # Creates a new instance of the Document class (from python-docx), i.e. a
    # completely empty Word document, held in memory and not yet
    # saved to disk, to which content will gradually be added.

    # Separator blank line.

    doc.add_heading(f"Latest {TOPIC} News", level=0)
    # Calls the "add_heading" method on the document object to add a
    # heading with the text "Latest AI News"; the "level=0" parameter indicates that
    # it is the highest-level heading (equivalent, in a Word
    # document, to the main "Title", distinct from level 1, 2,
    # etc. subheadings).
    stamp = doc.add_paragraph(
        f"Neutral cross-source summary — generated {datetime.now():%Y-%m-%d %H:%M}"
    )
    # Adds a new normal text paragraph to the document, built
    # as an f-string: the fixed text "Neutral cross-source summary —
    # generated " is followed by the current date and time, this time formatted
    # with both date and time (":%Y-%m-%d %H:%M"). The "paragraph" object
    # returned by add_paragraph is saved into the variable
    # "stamp", so it can be further modified in the next line.
    stamp.runs[0].italic = True
    # In python-docx, the text of a paragraph is internally organized into
    # a list of "runs" (blocks of text with uniform formatting,
    # accessible via the ".runs" attribute); here the first (and in
    # this case only) run of the just-created paragraph is accessed, index [0], and
    # its ".italic" property is set to True, making that text
    # appear in italics in the final Word document.

    # Separator blank line.

    for item in items:
        # "for" loop: iterates over each "item" dictionary contained in the
        # "items" list received as a parameter, i.e. over each
        # news item to insert into the document, one at a time and in the order in which
        # they appear in the list.
        heading = doc.add_heading(level=2)
        # Adds a new empty level-2 heading to the document (a
        # subheading, smaller than the main title but larger than
        # normal text); the created heading object is saved into the
        # local variable "heading" so its text can be set right
        # after (instead of passing the text directly to add_heading, so
        # that formatting can be controlled more finely via a
        # "run", as in the following line).
        heading.add_run(f"{item['title']}:")
        # Adds a text "run" to the just-created heading, i.e. a
        # segment of text with uniform formatting, whose content is
        # the value associated with the "title" key of the current
        # "item" dictionary (dictionary access with the item['title'] syntax),
        # followed by a literal ":" character, according to the format
        # required by the file's initial docstring ("news title:").

        # Separator blank line.

        para = doc.add_paragraph(item["explanation"])
        # Adds a new normal text paragraph to the document, whose
        # content is the value associated with the "explanation" key of the
        # current "item" dictionary (the neutral explanation of the news item);
        # the resulting paragraph object is saved into the
        # local variable "para" so its formatting can be modified right after.
        para.paragraph_format.left_indent = Pt(24)
        # Accesses the paragraph's formatting properties via
        # the ".paragraph_format" attribute and sets the left indent
        # ("left_indent", i.e. how much the text is shifted to the right
        # relative to the margin) to 24 typographic points (created with the
        # "Pt" class imported at the beginning of the file), so that the explanation
        # appears visually "indented" below its own
        # heading, reproducing the format described in the initial docstring.
        para.paragraph_format.space_after = Pt(10)
        # Sets, again via the paragraph's formatting properties,
        # the vertical space to leave AFTER the paragraph
        # itself (before the following content) to 10 typographic points,
        # to give a more pleasant visual separation between one news item and
        # the next.

    # Separator blank line (outside the for loop over the news items).

    if INCLUDE_SOURCES and sources:
        # Checks, with the "and" operator, two conditions: that the
        # global constant INCLUDE_SOURCES is True (the user wants to include the
        # sources section) AND that the "sources" list received as a
        # parameter is not empty (i.e. that there is actually at least one
        # source to show); only if both are true is the
        # indented block below executed.
        doc.add_page_break()
        # Inserts an explicit page break into the document,
        # making everything that follows (the "Sources" section) start
        # on a new page of the Word document, separating it visually
        # from the list of news items.
        doc.add_heading("Sources", level=1)
        # Adds a heading with the text "Sources", this time at level
        # 1 (a section heading, larger than a level-2 subheading
        # but smaller than the main level-0 title).
        for title, url in sources:
            # "for" loop with tuple "unpacking": at each
            # iteration, the current element of the "sources" list (a
            # two-element tuple) is automatically unpacked into the
            # two local variables "title" and "url".
            p = doc.add_paragraph(style="List Bullet")
            # Adds a new empty paragraph to the document, applying Word's
            # built-in "List Bullet" style (bulleted list), which
            # will make a bullet point appear before the text that will
            # be added to this paragraph. The paragraph object is
            # saved into the local variable "p".
            p.add_run(f"{title} — ").bold = True
            # Adds a first text "run" to the "p" paragraph, containing
            # the source title followed by a space, an em dash
            # "—" and another space (f-string); the expression returned by
            # ".add_run(...)" is the just-created run object, on which
            # the ".bold" property is immediately set to True, making
            # that portion of text bold in the final document.
            p.add_run(url)
            # Adds a second text "run" to the same "p" paragraph,
            # this time containing simply the source's URL, without
            # further formatting (not bold), which will therefore appear
            # right after the bold title, on the same line.

    # Separator blank line.

    doc.save(path)
    # Calls the document object's "save" method to physically write
    # to disk, at the path indicated by the "path" string (received as
    # a parameter of the function), all the content added so far in
    # memory, producing an actual .docx file that can be opened with Microsoft Word or
    # compatible programs.


# -----------------------------------------------------------------------------
# Main
# -----------------------------------------------------------------------------
# Section comment: introduces the main function that orchestrates
# the entire program flow, from start to finish.

def main() -> None:
    # Defines the "main" function, with no parameters, which does not return
    # any useful value ("-> None" annotation): it represents the logical
    # entry point of the entire script, which calls in sequence all the
    # other functions defined above.
    client = make_client()
    # Calls the previously defined "make_client" function to obtain
    # a client object already configured and authenticated against Vertex AI,
    # saving it into the local variable "client".

    # Separator blank line.

    print(f"[info] gathering {TOPIC} news (grounded with Google Search)...",
          file=sys.stderr)
    # Prints an informational message to stderr signaling to the user that
    # the program is about to start gathering news via
    # web search, to give visual feedback while waiting (the
    # call to the model can take a few seconds).
    resp, _ = generate(client, build_prompt(), use_search=True)
    # Calls the "generate" function defined above, passing the client, the
    # prompt text (obtained by calling "build_prompt()") and the keyword
    # parameter "use_search=True" to enable Google search. The function
    # returns a two-element tuple; the first (the actual
    # response) is assigned to the "resp" variable, while the second (the
    # name of the model used) is assigned to the special name "_", which by
    # convention in Python indicates a value that is received but not of interest
    # to use afterward.

    # Separator blank line.

    try:
        # Opens a "try" block because interpreting the response
        # text as JSON (in the following line) could fail if the
        # model returned text that does not conform to expectations.
        items = parse_items(resp.text)
        # Calls the "parse_items" function defined above, passing the
        # raw text of the model's response ("resp.text"), and saves the
        # result (the list of well-formed title/explanation dictionaries)
        # into the local variable "items".
    except (json.JSONDecodeError, KeyError, TypeError) as exc:
        # Catches three possible categories of error during parsing:
        # "json.JSONDecodeError" if the text is not valid JSON,
        # "KeyError" if an expected dictionary didn't contain the
        # required keys (title/explanation), "TypeError" if a value had a
        # type incompatible with the required operations; the caught
        # exception is assigned to the local name "exc".
        print(f"[error] could not parse model output as JSON: {exc}",
              file=sys.stderr)
        # Prints an error message to stderr explaining that it wasn't
        # possible to interpret the model's output as JSON, including
        # the detail of the caught exception.
        print("----- raw output -----\n" + resp.text, file=sys.stderr)
        # Prints to stderr, for diagnostic purposes, the entire raw text
        # returned by the model (concatenated with the "+" operator to
        # a descriptive heading and a newline "\n"), so that
        # the user can manually inspect what the model actually
        # generated and understand why parsing failed.
        sys.exit(1)
        # Immediately terminates execution of the entire program with
        # exit code 1, a value conventionally used by operating
        # systems to indicate that the program terminated due to an
        # error (0 would instead indicate termination without errors).

    # Separator blank line (reached only if parsing succeeded).

    if not items:
        # Checks whether the "items" list is empty: the "not" operator applied
        # to a list returns True if the list contains no elements
        # (an empty list is considered "false" in a boolean context in
        # Python).
        print("[error] no news items were produced.", file=sys.stderr)
        # If no valid news item was produced (even if perhaps the
        # JSON was syntactically correct but empty, or all elements
        # were discarded by the filter in parse_items), prints to stderr
        # an error message reporting this.
        sys.exit(1)
        # Terminates the program with exit code 1, indicating an error,
        # for the same reason explained above.

    # Separator blank line.

    sources = extract_sources(resp)
    # Calls the "extract_sources" function defined above, passing
    # the complete response object "resp" (which also contains the
    # grounding metadata), and saves the result (the list of unique
    # title/url source pairs) into the local variable "sources".
    print(f"[info] {len(sources)} source links used by the model:", file=sys.stderr)
    # Prints an informational message to stderr indicating how many sources
    # (computed with the built-in function "len", which returns the number
    # of elements in a list) were actually used by the model.
    for title, url in sources:
        # "for" loop with tuple unpacking, similar to the one seen
        # in write_docx: iterates over each (title, url) pair of the
        # "sources" list.
        print(f"    - {title}: {url}", file=sys.stderr)
        # For each source, prints a line to stderr formatted with
        # four spaces of indent, a dash, the title, a colon and
        # the URL, to show the user a readable list of the sources
        # directly in the terminal, in addition to those that will be written
        # into the Word document.

    # Separator blank line.

    write_docx(items, sources, OUTPUT_FILE)
    # Calls the "write_docx" function defined above, passing the list
    # of news items, the list of sources and the output path computed
    # at the beginning of the file (the global constant OUTPUT_FILE), to physically
    # generate the final Word document on disk.

    # Separator blank line.

    print(f"[done] wrote {len(items)} items to {OUTPUT_FILE}")
    # Prints (this time to standard stdout, not stderr, since the
    # "file" parameter is not specified, which defaults to sys.stdout) a
    # completion message confirming how many news items were
    # written and to which file, representing the "final",
    # user-facing output of the program.
    if INCLUDE_SOURCES:
        # Checks the global constant INCLUDE_SOURCES again: if true,
        # also prints a message about the included sources.
        print(f"[done] included {len(sources)} sources")
        # Prints a message to stdout confirming how many sources were
        # included in the final document.


if __name__ == "__main__":
    # This is a standard Python idiom: the special variable "__name__"
    # automatically equals the string "__main__" when the file is
    # executed directly (for example with "python ai_news_report.py"),
    # while it takes the module's own name if the file is instead
    # imported by another Python script (with "import
    # ai_news_mainScript_adcGcloud"). This check therefore allows
    # the following block to be executed only when the script is launched
    # directly, and not when its functions are reused elsewhere
    # via import.
    main()
    # Calls the "main" function defined above, actually starting
    # the entire program flow: client creation, prompt generation,
    # call to the model, response parsing, source extraction and
    # writing of the final Word document.
